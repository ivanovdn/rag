import re
import uuid
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from slugify import slugify

from config import settings
from ingest.chunk_models import PolicyChunk
from ingest.numbering import NumberingResolver


def extract_heading_level(paragraph) -> int | None:
    """Returns 1, 2, 3 etc. for Heading styles, None otherwise."""
    style_name = paragraph.style.name or ""
    if style_name.startswith("Heading"):
        try:
            return int(style_name.split()[-1])
        except (ValueError, IndexError):
            return None
    return None


def extract_clause_number(text: str) -> str | None:
    """Regex: matches 1.2, 4.2.1, 3. etc at start of paragraph."""
    pattern = r"^(\d+(?:\.\d+)*\.?)\s"
    match = re.match(pattern, text.strip())
    return match.group(1).rstrip(".") if match else None


def extract_label(text: str) -> str | None:
    """Extract 'Label:' prefix from List Paragraph items (e.g. 'Acceptable Use: ...')."""
    match = re.match(r"^([A-Z][A-Za-z\s&\-/()]+):\s", text.strip())
    return match.group(1).strip() if match else None


def extract_clause_name(para) -> str:
    """
    Extract clause name from the bold first run of a paragraph.

    ilvl=1 paragraphs look like:
        runs = [bold:"Blogging and Social Media:", normal:" Limited and occasional..."]

    Returns "Blogging and Social Media" (bold text, colon stripped).
    Returns empty string if first run is not bold.
    """
    runs = para.runs
    if not runs or not runs[0].bold:
        return ""

    # Collect all consecutive bold runs (sometimes name spans multiple runs)
    bold_parts = []
    for run in runs:
        if run.bold:
            bold_parts.append(run.text)
        else:
            break

    name = "".join(bold_parts).strip().rstrip(":").strip()
    return name


def _is_bold_heading(para, num_info) -> bool:
    """
    Detect bold-as-heading pattern: short, all-bold Normal paragraphs
    that serve as section headings in documents without Heading styles.
    """
    if not para.style.name.startswith("Normal"):
        return False

    # Skip if has Word auto-numbering (handled by NumberingResolver)
    if num_info and num_info["numFmt"] == "decimal":
        return False

    text = para.text.strip()
    if not text:
        return False

    # Must have runs, all bold
    text_runs = [r for r in para.runs if r.text.strip()]
    if not text_runs:
        return False
    if not all(r.bold for r in text_runs):
        return False

    # Short text only (headings are brief)
    if len(text.split()) > 10:
        return False

    # Skip metadata lines
    skip_prefixes = {"version", "date", "created by", "approved by", "managed by", "sensitivity"}
    if any(text.lower().startswith(p) for p in skip_prefixes):
        return False

    # Skip if it looks like a clause number (e.g., "4.2 Something")
    if re.match(r"^\d+(\.\d+)*\.?\s", text):
        return False

    return True


def _estimate_tokens(text: str) -> int:
    """Rough token estimate: ~4 chars per token."""
    return len(text) // 4


def _split_oversized(text: str, max_tokens: int) -> list[str]:
    """Split text at sentence boundaries if it exceeds max_tokens."""
    if _estimate_tokens(text) <= max_tokens:
        return [text]

    sentences = re.split(r"(?<=[.!?])\s+", text)
    parts = []
    current = []
    current_len = 0

    for sentence in sentences:
        s_tokens = _estimate_tokens(sentence)
        if current and current_len + s_tokens > max_tokens:
            parts.append(" ".join(current))
            current = [sentence]
            current_len = s_tokens
        else:
            current.append(sentence)
            current_len += s_tokens

    if current:
        parts.append(" ".join(current))

    return parts


def _table_to_text(table) -> str:
    """Convert a DOCX table to text rows: 'Header1: Val1 | Header2: Val2'."""
    rows = table.rows
    if not rows:
        return ""

    headers = [cell.text.strip() for cell in rows[0].cells]
    lines = []

    for row in rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if headers and any(h for h in headers):
            pairs = [f"{h}: {c}" for h, c in zip(headers, cells)]
            lines.append(" | ".join(pairs))
        else:
            lines.append(" | ".join(cells))

    if not lines and headers:
        lines.append(" | ".join(headers))

    return "\n".join(lines)


def _make_chunk(
    text: str,
    doc_id: str,
    doc_title: str,
    doc_filename: str,
    doc_link: str,
    section: str,
    section_number: str,
    clause: str,
    clause_number: str,
    chunk_index: int,
    last_updated: str,
) -> PolicyChunk:
    # Build display string
    parts = []
    if section_number and section:
        parts.append(f"{section_number}. {section}")
    elif section:
        parts.append(section)
    if clause_number and clause:
        parts.append(f"{clause_number}. {clause}")
    elif clause:
        parts.append(clause)
    section_display = " > ".join(parts)

    return PolicyChunk(
        chunk_id=str(uuid.uuid4()),
        doc_id=doc_id,
        doc_title=doc_title,
        doc_filename=doc_filename,
        doc_link=doc_link,
        section=section,
        section_number=section_number,
        clause=clause,
        clause_number=clause_number,
        section_display=section_display,
        text=text.strip(),
        char_count=len(text.strip()),
        chunk_index=chunk_index,
        last_updated=last_updated,
    )


def parse_docx(filepath: Path, doc_link: str) -> list[PolicyChunk]:
    """Main parser. Returns flat list of PolicyChunk objects."""
    doc = Document(filepath)
    resolver = NumberingResolver(doc)

    doc_filename = filepath.name
    doc_title = filepath.stem.replace("_", " ").title()
    doc_id = slugify(filepath.stem)
    last_updated = datetime.now(timezone.utc).isoformat()

    chunks: list[PolicyChunk] = []
    current_section = ""
    current_section_number = ""
    current_clause = ""
    current_clause_number = ""
    current_text_buffer: list[str] = []
    # Accumulator for undersized chunks within the same section
    pending_small: list[str] = []
    pending_clause = ""
    chunk_index = 0

    def _emit_chunk(text: str, clause_num: str):
        """Create chunk(s) from text, splitting if oversized."""
        nonlocal chunk_index
        parts = _split_oversized(text, settings.chunk_max_tokens)
        for part in parts:
            c = _make_chunk(
                text=part,
                doc_id=doc_id,
                doc_title=doc_title,
                doc_filename=doc_filename,
                doc_link=doc_link,
                section=current_section,
                section_number=current_section_number,
                clause=current_clause,
                clause_number=current_clause_number or clause_num,
                chunk_index=chunk_index,
                last_updated=last_updated,
            )
            chunks.append(c)
            chunk_index += 1

    def flush_buffer():
        """Flush current buffer: emit if large enough, otherwise accumulate."""
        nonlocal current_text_buffer, current_clause_number, pending_small, pending_clause
        text = "\n".join(current_text_buffer).strip()
        current_text_buffer = []
        if not text:
            return

        if _estimate_tokens(text) >= settings.chunk_min_tokens:
            # First emit any pending small text
            flush_pending()
            _emit_chunk(text, current_clause_number)
        else:
            # Accumulate into pending
            pending_small.append(text)
            if not pending_clause and current_clause_number:
                pending_clause = current_clause_number

    def flush_pending():
        """Emit accumulated small chunks as one combined chunk."""
        nonlocal pending_small, pending_clause
        if not pending_small:
            return
        combined = "\n".join(pending_small).strip()
        pending_small = []
        if combined:
            _emit_chunk(combined, pending_clause)
        pending_clause = ""

    def flush_all():
        """Flush both buffer and pending at section/document boundaries."""
        flush_buffer()
        flush_pending()

    # Build a lookup from element to paragraph object for fast access
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    # Iterate doc.element.body to handle paragraphs and tables in order
    for element in doc.element.body:
        if element.tag == qn("w:p"):
            para = para_map.get(element)
            if para is None:
                continue

            text = para.text.strip()
            if not text:
                continue

            # -- Resolve Word auto-numbering --
            num_info = resolver.resolve(element)
            level = extract_heading_level(para)

            if num_info and num_info["numFmt"] == "decimal":
                resolved = num_info["resolved"]  # e.g. "7.5."
                # Strip trailing dot for clean number
                clean_number = resolved.rstrip(".")  # "7" or "7.5"

                # Prepend number to text for the chunk content
                text = f"{resolved} {text}"

                # Use number depth to determine level (more reliable than ilvl)
                # "4" → depth 1 (section), "4.2" → depth 2 (clause), "4.2.1" → depth 3+
                num_depth = len(clean_number.split("."))

                if num_depth == 1:
                    # Section level
                    flush_all()
                    # Use bold text as section name if available (cleaner),
                    # otherwise fall back to full paragraph text
                    section_name = extract_clause_name(para)  # gets bold prefix
                    if not section_name:
                        # Truncate long section names to first phrase
                        raw = para.text.strip()
                        if len(raw.split()) > 8:
                            for sep in [":", ",", "."]:
                                if sep in raw:
                                    raw = raw[:raw.index(sep)].strip()
                                    break
                        section_name = raw
                    current_section = section_name
                    current_section_number = clean_number
                    current_clause = ""
                    current_clause_number = ""
                    level = None  # already handled, skip the level block below

                elif num_depth >= 2:
                    # Clause level
                    flush_all()
                    current_clause = extract_clause_name(para) or para.text.strip()
                    current_clause_number = clean_number
                    # The text goes into the buffer as chunk content
                    current_text_buffer.append(text)
                    level = None  # skip heading block

            if level:
                # Standard heading processing (for docs without numbering)
                flush_all()
                idx = min(level, 3) - 1
                if idx == 0:
                    current_section = text
                    current_section_number = ""
                    current_clause = ""
                    current_clause_number = ""
                elif idx == 1:
                    current_clause = text
                    current_clause_number = ""

            elif _is_bold_heading(para, num_info):
                # Bold-as-heading: treat as section heading
                flush_all()
                current_section = text
                current_section_number = ""
                current_clause = ""
                current_clause_number = ""

            elif num_info is None or num_info["numFmt"] != "decimal":
                # Regular paragraph or bullet — clause detection logic
                clause_num = extract_clause_number(text)
                if not clause_num:
                    label = extract_label(text)
                    if label and para.style.name == "List Paragraph":
                        # Label goes to clause NAME, not clause_number
                        flush_buffer()
                        current_clause = label
                        current_clause_number = ""
                        current_text_buffer.append(text)
                        continue
                if clause_num and current_text_buffer:
                    flush_buffer()
                    current_clause_number = clause_num
                elif clause_num:
                    current_clause_number = clause_num
                current_text_buffer.append(text)

        elif element.tag == qn("w:tbl"):
            table = table_map.get(element)
            if table:
                table_text = _table_to_text(table)
                if table_text:
                    current_text_buffer.append(table_text)

    flush_all()

    # ── Filter noise chunks ──
    NOISE_SECTIONS = {
        "revision history",
        "change history",
        "table of contents",
        "document control",
        "version control",
        "approval history",
    }

    filtered = []
    for c in chunks:
        section_lower = (c.section or "").lower().strip()

        # Skip noise sections
        if section_lower in NOISE_SECTIONS:
            continue

        # Skip title-only chunks (document name with no real content)
        if not c.section and not c.clause and _estimate_tokens(c.text) < 15:
            continue

        # Skip near-empty chunks
        if len(c.text.strip()) < 20:
            continue

        filtered.append(c)

    # Re-index
    for i, c in enumerate(filtered):
        c.chunk_index = i

    return filtered
